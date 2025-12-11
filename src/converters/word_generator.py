#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档生成器模块
负责将解析后的Markdown内容转换为Word文档
"""

import os
import re
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
from dataclasses import dataclass

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    raise ImportError("请安装python-docx库: pip install python-docx")

from .markdown_parser import MarkdownElement

# 图表相关模块（可选导入）
try:
    # 尝试相对导入
    try:
        from ..utils.chart_recognizer import ChartRecognizer
        from ..utils.chart_generator import ChartGenerator
    except ImportError:
        # 如果相对导入失败，尝试绝对导入
        import sys
        import os
        # 获取当前文件所在目录的父目录（src目录）
        current_dir = os.path.dirname(os.path.abspath(__file__))
        src_dir = os.path.dirname(current_dir)
        if src_dir not in sys.path:
            sys.path.insert(0, src_dir)
        from utils.chart_recognizer import ChartRecognizer
        from utils.chart_generator import ChartGenerator
    CHARTS_AVAILABLE = True
except ImportError as e:
    CHARTS_AVAILABLE = False
    print(f"图表模块导入失败: {e}")


@dataclass
class WordStyle:
    """Word样式配置"""
    font_name: str = "微软雅黑"
    font_size: int = 12
    line_spacing: float = 1.15
    paragraph_spacing_before: int = 0
    paragraph_spacing_after: int = 6
    alignment: str = "left"
    bold: bool = False
    italic: bool = False
    color: Optional[str] = None


class WordGenerator:
    """Word文档生成器"""
    
    def __init__(self, template_path: Optional[str] = None, config: Optional[Dict] = None, style_handler=None):
        """初始化生成器
        
        Args:
            template_path: Word模板文件路径
            config: 生成器配置
            style_handler: 样式处理器（可选）
        """
        self.config = config or {}
        self.template_path = template_path
        self.style_handler = style_handler
        self.document = self._create_document()
        self.styles = self._setup_styles()
        
        # 图表相关配置
        self.enable_charts = self.config.get('enable_charts', False)
        self.chart_data = []  # 存储识别的图表数据
        self.chart_images = {}  # 存储生成的图片路径 {position: image_path}
        self.chart_generator = None
        self.temp_image_files = []  # 临时图片文件列表，用于清理
        self.chart_data_source = self.config.get('chart_data', '')  # 图表数据源
        
    def _create_document(self) -> Document:
        """创建Word文档"""
        if self.template_path and os.path.exists(self.template_path):
            return Document(self.template_path)
        else:
            return Document()
    
    def _setup_styles(self) -> Dict[str, WordStyle]:
        """设置文档样式"""
        styles = {
            'normal': WordStyle(
                font_name="微软雅黑",
                font_size=12,
                line_spacing=1.15
            ),
            'heading1': WordStyle(
                font_name="微软雅黑",
                font_size=18,
                bold=True,
                paragraph_spacing_before=12,
                paragraph_spacing_after=6
            ),
            'heading2': WordStyle(
                font_name="微软雅黑",
                font_size=16,
                bold=True,
                paragraph_spacing_before=10,
                paragraph_spacing_after=6
            ),
            'heading3': WordStyle(
                font_name="微软雅黑",
                font_size=14,
                bold=True,
                paragraph_spacing_before=8,
                paragraph_spacing_after=6
            ),
            'heading4': WordStyle(
                font_name="微软雅黑",
                font_size=13,
                bold=True,
                paragraph_spacing_before=6,
                paragraph_spacing_after=6
            ),
            'heading5': WordStyle(
                font_name="微软雅黑",
                font_size=12,
                bold=True,
                paragraph_spacing_before=6,
                paragraph_spacing_after=6
            ),
            'heading6': WordStyle(
                font_name="微软雅黑",
                font_size=12,
                bold=True,
                italic=True,
                paragraph_spacing_before=6,
                paragraph_spacing_after=6
            ),
            'code': WordStyle(
                font_name="Consolas",
                font_size=10,
                color="#333333"
            ),
            'quote': WordStyle(
                font_name="微软雅黑",
                font_size=12,
                italic=True,
                color="#666666"
            )
        }
        
        # 应用自定义样式配置
        if 'styles' in self.config:
            for style_name, style_config in self.config['styles'].items():
                if style_name in styles:
                    for attr, value in style_config.items():
                        setattr(styles[style_name], attr, value)
        
        return styles
    
    def generate(self, markdown_element: MarkdownElement, output_path: str, markdown_text: Optional[str] = None) -> bool:
        """生成Word文档
        
        Args:
            markdown_element: 解析后的Markdown元素
            output_path: 输出文件路径
            markdown_text: 原始Markdown文本（用于图表识别）
            
        Returns:
            是否生成成功
        """
        try:
            # 如果启用图表功能，先识别和生成图表
            print(f"图表功能检查: enable_charts={self.enable_charts}, CHARTS_AVAILABLE={CHARTS_AVAILABLE}, markdown_text={'有' if markdown_text else '无'}")
            if self.enable_charts and CHARTS_AVAILABLE and markdown_text:
                print("开始准备图表...")
                self._prepare_charts(markdown_text)
            else:
                if not self.enable_charts:
                    print("图表功能已禁用")
                elif not CHARTS_AVAILABLE:
                    print("图表模块不可用")
                elif not markdown_text:
                    print("未提供markdown_text参数")
            
            # 处理文档内容
            self._process_element(markdown_element)
            
            # 处理未插入的图表（插入到文档末尾）
            if self.chart_images:
                print(f"文档处理完成，检查未插入的图表: 剩余 {len(self.chart_images)} 个图表")
            self._insert_remaining_charts()
            
            # 保存文档
            self.document.save(output_path)
            
            # 清理临时图片文件
            self._cleanup_chart_images()
            
            return True
            
        except Exception as e:
            print(f"生成Word文档失败: {e}")
            # 确保清理临时文件
            self._cleanup_chart_images()
            return False
    
    def generate_from_html(self, html_content: str, metadata: Dict[str, Any], output_path: str) -> bool:
        """从HTML内容生成Word文档（简化版本）
        
        Args:
            html_content: HTML内容
            metadata: 文档元数据
            output_path: 输出文件路径
            
        Returns:
            是否生成成功
        """
        try:
            from bs4 import BeautifulSoup
            
            # 解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 添加文档标题
            if metadata.get('title'):
                self._add_title(metadata['title'])
            
            # 处理HTML元素
            self._process_html_elements(soup)
            
            # 保存文档
            self.document.save(output_path)
            return True
            
        except Exception as e:
            print(f"生成Word文档失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _process_html_elements(self, soup):
        """处理HTML元素"""
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'blockquote', 'ul', 'ol', 'table']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                self.document.add_heading(element.get_text().strip(), level=level)
            elif element.name == 'p':
                paragraph = self.document.add_paragraph()
                self._process_paragraph_content(paragraph, element)
            elif element.name == 'pre':
                code_text = element.get_text().strip()
                code_paragraph = self.document.add_paragraph(code_text)
                self._apply_style(code_paragraph, self.styles['code'])
            elif element.name == 'blockquote':
                quote_text = element.get_text().strip()
                quote_paragraph = self.document.add_paragraph(quote_text)
                self._apply_style(quote_paragraph, self.styles['quote'])
                quote_paragraph.paragraph_format.left_indent = Inches(0.5)
            elif element.name in ['ul', 'ol']:
                self._process_html_list(element)
            elif element.name == 'table':
                # 在表格前添加一个空行
                self.document.add_paragraph()
                self._process_html_table(element)
                # 在表格后添加一个空行
                self.document.add_paragraph()
    
    def _process_paragraph_content(self, paragraph, element):
        """处理段落内容"""
        for content in element.contents:
            if hasattr(content, 'name'):
                if content.name == 'strong' or content.name == 'b':
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                elif content.name == 'em' or content.name == 'i':
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                elif content.name == 'code':
                    run = paragraph.add_run(content.get_text())
                    self._apply_code_style(run)
                elif content.name == 'a':
                    text = content.get_text()
                    url = content.get('href', '#')
                    self._add_hyperlink(paragraph, url, text)
                else:
                    paragraph.add_run(content.get_text())
            else:
                paragraph.add_run(str(content))
    
    def _process_html_list(self, element, indent_level: int = 0):
        """处理HTML列表
        
        Args:
            element: HTML列表元素
            indent_level: 缩进级别（用于嵌套列表）
        """
        is_ordered = element.name == 'ol'
        
        # 计算左缩进：基础缩进 + 嵌套缩进
        # 基础缩进：0.5 英寸（约 1.27 厘米），使列表比段落有更多缩进
        # 嵌套缩进：每级嵌套增加 0.5 英寸
        base_indent = Inches(0.5)
        nested_indent = Inches(0.5 * indent_level)
        total_indent = base_indent + nested_indent
        
        for li in element.find_all('li', recursive=False):
            # 根据列表类型选择样式
            if is_ordered:
                # 有序列表使用 List Number 样式
                paragraph = self.document.add_paragraph(style='List Number')
            else:
                # 无序列表使用 List Bullet 样式
                paragraph = self.document.add_paragraph(style='List Bullet')
            
            # 设置列表项的左缩进，使其比普通段落有更多缩进
            paragraph.paragraph_format.left_indent = total_indent
            
            # 处理列表项内容，包括格式化文本（粗体、斜体、代码、链接等）
            self._process_paragraph_content(paragraph, li)
            
            # 处理嵌套列表
            nested_lists = li.find_all(['ul', 'ol'], recursive=False)
            for nested_list in nested_lists:
                self._process_html_list(nested_list, indent_level + 1)
    
    def _set_cell_vertical_alignment(self, cell, alignment='center'):
        """设置单元格垂直对齐方式
        
        Args:
            cell: Word表格单元格对象
            alignment: 对齐方式，可选值：'top', 'center', 'bottom'，默认为'center'
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # 查找或创建vAlign元素
        vAlign_elements = tcPr.findall(qn('w:vAlign'))
        if vAlign_elements:
            vAlign = vAlign_elements[0]
        else:
            vAlign = OxmlElement('w:vAlign')
            tcPr.append(vAlign)
        
        # 设置垂直对齐方式
        vAlign.set(qn('w:val'), alignment)
    
    def _set_cell_background(self, cell, rgb_color):
        """设置单元格背景色
        
        Args:
            cell: Word表格单元格对象
            rgb_color: RGB颜色元组，如 (141, 179, 226)
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # 查找或创建shd元素
        shd_elements = tcPr.findall(qn('w:shd'))
        if shd_elements:
            shd = shd_elements[0]
        else:
            shd = OxmlElement('w:shd')
            tcPr.append(shd)
        
        # 将RGB转换为十六进制（去掉#号）
        hex_color = f"{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}"
        shd.set(qn('w:fill'), hex_color)
    
    def _set_table_header_style(self, table):
        """设置表格表头样式（背景色）
        
        Args:
            table: Word表格对象
        """
        if table.rows is None or len(table.rows) == 0:
            return
        
        # 表头背景色 RGB(141, 179, 226)
        header_color = (141, 179, 226)
        
        # 设置第一行（表头）的背景色
        header_row = table.rows[0]
        for cell in header_row.cells:
            self._set_cell_background(cell, header_color)
    
    def _merge_first_column_cells(self, table):
        """合并第一列中相同值的单元格，并设置垂直居中
        
        Args:
            table: Word表格对象
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if table.rows is None or len(table.rows) == 0:
            return
        
        num_rows = len(table.rows)
        if num_rows <= 1:
            return
        
        # 收集第一列的值，用于判断哪些行需要合并
        first_col_values = []
        for i in range(num_rows):
            if len(table.rows[i].cells) > 0:
                first_col_values.append(table.rows[i].cells[0].text.strip())
            else:
                first_col_values.append("")
        
        # 合并第一列中相同值的连续单元格（从第二行开始，跳过表头）
        i = 1  # 从第二行开始，跳过表头
        while i < num_rows:
            current_value = first_col_values[i]
            if not current_value:  # 跳过空值
                i += 1
                continue
            
            # 找到相同值的连续范围
            start_row = i
            end_row = i
            
            # 向后查找相同值的行
            j = i + 1
            while j < num_rows:
                if first_col_values[j] == current_value:
                    end_row = j
                    j += 1
                else:
                    break
            
            # 如果找到需要合并的行（至少2行）
            if end_row > start_row:
                # 保存第一个单元格的值（合并前）
                start_cell = table.rows[start_row].cells[0]
                cell_value = current_value  # 使用收集到的值，避免重复
                
                # 合并单元格
                end_cell = table.rows[end_row].cells[0]
                start_cell.merge(end_cell)
                
                # 清空合并后的单元格内容，然后重新设置值（避免重复）
                start_cell.text = ""
                start_cell.text = cell_value
                
                # 设置垂直居中对齐
                self._set_cell_vertical_alignment(start_cell, 'center')
                
                # 确保合并后的单元格内容水平居中，并设置字体样式
                for paragraph in start_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 设置字体为微软雅黑，小五（9号）
                    for run in paragraph.runs:
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(9)
            
            # 移动到下一组
            i = end_row + 1
    
    def _process_html_table(self, element):
        """处理HTML表格"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        rows = element.find_all('tr')
        if not rows:
            return
        
        # 计算列数
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        
        # 创建表格
        table = self.document.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        # 禁用自动调整列宽，以便手动控制列宽
        table.autofit = False
        
        # 设置表格属性：对齐方式、文字环绕、单元格边距
        self._setup_table_properties(table)
        
        # 计算每列的最大内容长度（用于智能列宽分配）
        column_max_lengths = [0] * max_cols
        for row in rows:
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    cell_text = cell.get_text().strip()
                    # 计算文本长度（中文字符按2个字符计算）
                    text_length = sum(2 if ord(c) > 127 else 1 for c in cell_text)
                    column_max_lengths[j] = max(column_max_lengths[j], text_length)
        
        # 填充数据并设置样式
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    word_cell = table.cell(i, j)
                    word_cell.text = cell.get_text().strip()
                    
                    # 设置单元格水平对齐方式为居中
                    for paragraph in word_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 设置字体为微软雅黑，小五（9号）
                        for run in paragraph.runs:
                            run.font.name = '微软雅黑'
                            run.font.size = Pt(9)
                    
                    # 设置单元格垂直对齐方式为居中
                    self._set_cell_vertical_alignment(word_cell, 'center')
        
        # 合并第一列中相同值的单元格
        self._merge_first_column_cells(table)
        
        # 设置表头背景色
        self._set_table_header_style(table)
        
        # 智能调整列宽（设置表格整体宽度最小值和最大值）
        self._adjust_table_column_widths(table, column_max_lengths, min_table_width_cm=8.0, max_table_width_cm=16.0)
    
    def _process_element(self, element: MarkdownElement):
        """处理Markdown元素"""
        if element.element_type == 'document':
            self._process_document(element)
        elif element.element_type.startswith('heading'):
            self._process_heading(element)
        elif element.element_type == 'paragraph':
            self._process_paragraph(element)
        elif element.element_type == 'code_block':
            self._process_code_block(element)
        elif element.element_type == 'table':
            self._process_table(element)
        elif element.element_type == 'image':
            self._process_image(element)
        elif element.element_type == 'list':
            self._process_list(element)
        elif element.element_type == 'quote':
            self._process_quote(element)
        
        # 处理子元素
        for child in element.children:
            self._process_element(child)
    
    def _process_document(self, element: MarkdownElement):
        """处理文档根元素"""
        # 不再自动添加标题，标题应该由 Markdown 中的 H1 标题处理
        # if 'title' in element.attributes:
        #     self._add_title(element.attributes['title'])
        pass
    
    def _process_heading(self, element: MarkdownElement):
        """处理标题"""
        level = int(element.element_type.replace('heading', ''))
        style_name = f'heading{level}'
        
        paragraph = self.document.add_heading(element.content, level=level)
        
        # 优先使用 StyleHandler 的样式
        if self.style_handler:
            style = self.style_handler.get_style(style_name)
            if style:
                self._apply_style_handler_style(paragraph, style, is_heading=True)
                return
        
        # 否则使用默认样式
        self._apply_style(paragraph, self.styles.get(style_name, self.styles['normal']))
    
    def _process_paragraph(self, element: MarkdownElement):
        """处理段落"""
        paragraph = self.document.add_paragraph()
        
        # 处理段落内容，包括格式化文本
        self._process_formatted_text(paragraph, element.content)
        
        # 优先使用 StyleHandler 的样式
        if self.style_handler:
            style = self.style_handler.get_style('normal')
            if style:
                self._apply_style_handler_style(paragraph, style, is_heading=False)
            else:
                # 否则使用默认样式
                self._apply_style(paragraph, self.styles['normal'])
        else:
            # 否则使用默认样式
            self._apply_style(paragraph, self.styles['normal'])
        
        # 检查是否需要在此段落后插入图表
        if self.enable_charts and self.chart_images:
            self._check_and_insert_chart(paragraph, element.content)
    
    def _process_formatted_text(self, paragraph, text: str):
        """处理格式化文本"""
        # 处理粗体、斜体、代码等格式
        parts = self._split_formatted_text(text)
        
        for part in parts:
            if part['type'] == 'bold':
                run = paragraph.add_run(part['text'])
                run.bold = True
            elif part['type'] == 'italic':
                run = paragraph.add_run(part['text'])
                run.italic = True
            elif part['type'] == 'code':
                run = paragraph.add_run(part['text'])
                self._apply_code_style(run)
            elif part['type'] == 'link':
                self._add_hyperlink(paragraph, part['url'], part['text'])
            else:
                paragraph.add_run(part['text'])
    
    def _split_formatted_text(self, text: str) -> List[Dict[str, str]]:
        """分割格式化文本"""
        if not text:
            # 如果文本为空，返回空列表（不添加任何内容）
            return []
        
        parts = []
        current_pos = 0
        
        # 匹配各种格式的正则表达式
        patterns = [
            (r'\*\*([^*]+)\*\*', 'bold'),
            (r'\*([^*]+)\*', 'italic'),
            (r'`([^`]+)`', 'code'),
            (r'\[([^\]]+)\]\(([^)]+)\)', 'link')
        ]
        
        for pattern, format_type in patterns:
            for match in re.finditer(pattern, text):
                # 添加匹配前的普通文本
                if match.start() > current_pos:
                    parts.append({
                        'type': 'normal',
                        'text': text[current_pos:match.start()]
                    })
                
                # 添加格式化文本
                if format_type == 'link':
                    parts.append({
                        'type': 'link',
                        'text': match.group(1),
                        'url': match.group(2)
                    })
                else:
                    parts.append({
                        'type': format_type,
                        'text': match.group(1)
                    })
                
                current_pos = match.end()
        
        # 添加剩余的普通文本
        if current_pos < len(text):
            parts.append({
                'type': 'normal',
                'text': text[current_pos:]
            })
        
        # 如果没有找到任何部分，返回整个文本作为普通文本
        if not parts:
            parts.append({
                'type': 'normal',
                'text': text
            })
        
        return parts
    
    def _process_code_block(self, element: MarkdownElement):
        """处理代码块"""
        # 添加代码块标题（如果有语言标识）
        if 'language' in element.attributes:
            lang_paragraph = self.document.add_paragraph(f"代码 ({element.attributes['language']})")
            lang_paragraph.style = 'Caption'
        
        # 添加代码内容
        code_paragraph = self.document.add_paragraph(element.content)
        self._apply_style(code_paragraph, self.styles['code'])
        
        # 设置代码块背景色
        self._set_paragraph_background(code_paragraph, '#f8f8f8')
    
    def _process_table(self, element: MarkdownElement):
        """处理表格"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if 'rows' not in element.attributes or 'cols' not in element.attributes:
            return
        
        # 在表格前添加一个空行
        self.document.add_paragraph()
        
        rows = element.attributes['rows']
        cols = element.attributes['cols']
        
        table = self.document.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # 禁用自动调整列宽，以便手动控制列宽
        table.autofit = False
        
        # 设置表格属性：对齐方式、文字环绕、单元格边距
        self._setup_table_properties(table)
        
        # 填充表格数据并设置样式
        if 'data' in element.attributes:
            data = element.attributes['data']
            
            # 计算每列的最大内容长度（用于智能列宽分配）
            column_max_lengths = [0] * cols
            for i, row_data in enumerate(data):
                if i < rows:
                    for j, cell_data in enumerate(row_data):
                        if j < cols:
                            cell_text = str(cell_data)
                            # 计算文本长度（中文字符按2个字符计算）
                            text_length = sum(2 if ord(c) > 127 else 1 for c in cell_text)
                            column_max_lengths[j] = max(column_max_lengths[j], text_length)
            
            # 填充数据并设置样式
            for i, row_data in enumerate(data):
                if i < rows:
                    for j, cell_data in enumerate(row_data):
                        if j < cols:
                            cell = table.cell(i, j)
                            cell.text = str(cell_data)
                            
                            # 设置单元格水平对齐方式为居中
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                # 设置字体为微软雅黑，小五（9号）
                                for run in paragraph.runs:
                                    run.font.name = '微软雅黑'
                                    run.font.size = Pt(9)
                            
                            # 设置单元格垂直对齐方式为居中
                            self._set_cell_vertical_alignment(cell, 'center')
            
            # 合并第一列中相同值的单元格
            self._merge_first_column_cells(table)
            
            # 设置表头背景色
            self._set_table_header_style(table)
            
            # 智能调整列宽（设置表格整体宽度最小值和最大值）
            self._adjust_table_column_widths(table, column_max_lengths, min_table_width_cm=8.0, max_table_width_cm=16.0)
        
        # 在表格后添加一个空行
        self.document.add_paragraph()
    
    def _setup_table_properties(self, table):
        """设置表格属性：对齐方式、文字环绕、单元格边距
        
        Args:
            table: Word表格对象
        """
        # 1. 设置表格对齐方式为居中
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 2. 设置文字环绕方式为环绕
        tbl = table._tbl  # 获取底层的XML元素
        
        # 查找或创建tblPr元素
        tblPr_elements = tbl.findall(qn('w:tblPr'))
        if not tblPr_elements:
            tblPr = OxmlElement('w:tblPr')
            # 将tblPr插入到tbl的第一个位置
            tbl.insert(0, tblPr)
        else:
            tblPr = tblPr_elements[0]
        
        # 查找或创建tblpPr元素
        tblpPr_elements = tblPr.findall(qn('w:tblpPr'))
        if not tblpPr_elements:
            tblpPr = OxmlElement('w:tblpPr')
            tblPr.append(tblpPr)
        else:
            tblpPr = tblpPr_elements[0]
        
        # 设置文字环绕为环绕
        tblpPr.set(qn('w:wrap'), 'around')
        
        # 3. 设置单元格边距：上下均为0.15cm
        # 查找或创建tblCellMar元素
        tblCellMar_elements = tblPr.findall(qn('w:tblCellMar'))
        if not tblCellMar_elements:
            tblCellMar = OxmlElement('w:tblCellMar')
            tblPr.append(tblCellMar)
        else:
            tblCellMar = tblCellMar_elements[0]
        
        # 清除现有的边距设置（如果存在）
        for margin_name in ['w:top', 'w:bottom', 'w:left', 'w:right']:
            existing = tblCellMar.findall(qn(margin_name))
            for elem in existing:
                tblCellMar.remove(elem)
        
        # 设置上下边距为0.15cm
        top_margin = OxmlElement('w:top')
        top_margin.set(qn('w:w'), str(int(0.15 * 567)))  # 0.15cm转换为twips (1cm = 567 twips)
        top_margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(top_margin)
        
        bottom_margin = OxmlElement('w:bottom')
        bottom_margin.set(qn('w:w'), str(int(0.15 * 567)))
        bottom_margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(bottom_margin)
        
        # 设置左右边距为默认值（如果需要）
        left_margin = OxmlElement('w:left')
        left_margin.set(qn('w:w'), str(int(0.19 * 567)))  # 默认0.19cm
        left_margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(left_margin)
        
        right_margin = OxmlElement('w:right')
        right_margin.set(qn('w:w'), str(int(0.19 * 567)))  # 默认0.19cm
        right_margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(right_margin)
    
    def _adjust_table_column_widths(self, table, column_max_lengths, min_table_width_cm=8.0, max_table_width_cm=16.0):
        """根据内容长度智能调整表格列宽，并设置表格整体宽度的最小值和最大值
        
        Args:
            table: Word表格对象
            column_max_lengths: 每列的最大内容长度列表
            min_table_width_cm: 表格整体最小宽度（厘米），默认8.0cm
            max_table_width_cm: 表格整体最大宽度（厘米），默认16.0cm
        """
        if not column_max_lengths or sum(column_max_lengths) == 0:
            # 如果没有内容，设置默认宽度
            default_width_cm = (min_table_width_cm + max_table_width_cm) / 2
            if len(table.columns) > 0:
                col_width_cm = default_width_cm / len(table.columns)
                col_width = Cm(col_width_cm)
            else:
                col_width = Cm(default_width_cm)
            for col in table.columns:
                col.width = col_width
            return
        
        # 设置最小和最大列宽（厘米）
        min_col_width_cm = 1.2  # 最小列宽1.2cm（用于序号等短内容）
        max_col_width_cm = 6.0  # 最大列宽6.0cm（用于长描述）
        
        # 计算所有列的最大内容长度的中位数作为基准
        valid_lengths = [length for length in column_max_lengths if length > 0]
        if not valid_lengths:
            # 如果没有有效长度，使用默认宽度
            default_width_cm = (min_col_width_cm + max_col_width_cm) / 2
            column_widths_cm = [default_width_cm] * len(column_max_lengths)
        else:
            # 计算中位数
            sorted_lengths = sorted(valid_lengths)
            median_length = sorted_lengths[len(sorted_lengths) // 2]
            
            # 如果中位数为0或太小，使用平均值作为基准
            if median_length < 3:
                median_length = sum(valid_lengths) / len(valid_lengths) if valid_lengths else 10
            
            # 根据中位数计算基准宽度（中位数对应的列宽设为中等宽度）
            # 假设中位数对应的列宽为3.0cm（表格总宽度在8-16cm之间，5列的话平均约3cm）
            base_width_cm = 3.0
            char_width_cm = base_width_cm / median_length  # 每个字符对应的宽度
            
            # 根据内容计算每列需要的宽度
            column_widths_cm = []
            for max_length in column_max_lengths:
                if max_length > 0:
                    # 对于非常短的内容（<=3个字符），使用最小宽度
                    if max_length <= 3:
                        col_width_cm = min_col_width_cm
                    else:
                        # 根据字符数与中位数的比例计算列宽
                        # 如果内容长度等于中位数，则宽度为base_width_cm
                        # 如果内容长度是中位数的2倍，则宽度约为base_width_cm的2倍（但不超过最大值）
                        col_width_cm = max_length * char_width_cm
                        # 确保在最小和最大宽度之间
                        col_width_cm = max(min_col_width_cm, min(col_width_cm, max_col_width_cm))
                else:
                    # 如果列没有内容，使用最小宽度
                    col_width_cm = min_col_width_cm
                column_widths_cm.append(col_width_cm)
        
        # 计算表格总宽度
        total_width_cm = sum(column_widths_cm)
        
        # 如果总宽度超出范围，按比例缩放所有列
        if total_width_cm < min_table_width_cm:
            # 如果总宽度小于最小值，按比例放大
            scale_factor = min_table_width_cm / total_width_cm
            column_widths_cm = [w * scale_factor for w in column_widths_cm]
            total_width_cm = min_table_width_cm
        elif total_width_cm > max_table_width_cm:
            # 如果总宽度大于最大值，按比例缩小
            scale_factor = max_table_width_cm / total_width_cm
            column_widths_cm = [w * scale_factor for w in column_widths_cm]
            total_width_cm = max_table_width_cm
        
        # 确保每列宽度在合理范围内（缩放后可能超出范围）
        for j, width_cm in enumerate(column_widths_cm):
            width_cm = max(min_col_width_cm, min(width_cm, max_col_width_cm))
            column_widths_cm[j] = width_cm
        
        # 应用列宽 - 同时设置列和单元格的宽度以确保生效
        for j, width_cm in enumerate(column_widths_cm):
            col_width = Cm(width_cm)
            # 设置列的宽度
            table.columns[j].width = col_width
            # 同时设置该列所有单元格的宽度，确保生效
            for cell in table.columns[j].cells:
                cell.width = col_width
    
    def _process_image(self, element: MarkdownElement):
        """处理图片"""
        if 'src' not in element.attributes:
            return
        
        src = element.attributes['src']
        alt = element.attributes.get('alt', '')
        
        try:
            # 如果是本地文件路径
            if os.path.exists(src):
                paragraph = self.document.add_paragraph()
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.add_picture(src, width=Inches(6))
                
                # 添加图片说明
                if alt:
                    caption = self.document.add_paragraph(f"图: {alt}")
                    caption.style = 'Caption'
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # 网络图片或无效路径，添加占位符
                placeholder = self.document.add_paragraph(f"[图片: {alt or src}]")
                placeholder.style = 'Caption'
                
        except Exception as e:
            print(f"处理图片失败 {src}: {e}")
            # 添加错误占位符
            error_placeholder = self.document.add_paragraph(f"[图片加载失败: {alt or src}]")
            error_placeholder.style = 'Caption'
    
    def _process_list(self, element: MarkdownElement, indent_level: int = 0):
        """处理列表
        
        Args:
            element: 列表元素
            indent_level: 缩进级别（用于嵌套列表）
        """
        list_type = element.attributes.get('type', 'unordered')
        
        # 计算左缩进：基础缩进 + 嵌套缩进
        # 基础缩进：0.5 英寸（约 1.27 厘米），使列表比段落有更多缩进
        # 嵌套缩进：每级嵌套增加 0.5 英寸
        base_indent = Inches(0.5)
        nested_indent = Inches(0.5 * indent_level)
        total_indent = base_indent + nested_indent
        
        for item in element.children:
            # 根据列表类型选择样式
            if list_type == 'ordered':
                # 有序列表使用 List Number 样式
                paragraph = self.document.add_paragraph(style='List Number')
            else:
                # 无序列表使用 List Bullet 样式
                paragraph = self.document.add_paragraph(style='List Bullet')
            
            # 设置列表项的左缩进，使其比普通段落有更多缩进
            paragraph.paragraph_format.left_indent = total_indent
            
            # 处理列表项内容，包括格式化文本（粗体、斜体、代码、链接等）
            if item.content:
                self._process_formatted_text(paragraph, item.content)
            
            # 处理列表项的子元素（如嵌套列表、段落等）
            for child in item.children:
                if child.element_type == 'list':
                    # 嵌套列表：递归处理，缩进级别+1
                    self._process_list(child, indent_level + 1)
                elif child.element_type == 'paragraph':
                    # 列表项内的段落：作为列表项的一部分处理
                    self._process_formatted_text(paragraph, child.content)
    
    def _process_quote(self, element: MarkdownElement):
        """处理引用"""
        paragraph = self.document.add_paragraph(element.content)
        self._apply_style(paragraph, self.styles['quote'])
        
        # 设置引用样式
        paragraph.paragraph_format.left_indent = Inches(0.5)
        self._set_paragraph_border(paragraph, 'left')
    
    def _apply_style_handler_style(self, paragraph, element_style, is_heading=False):
        """应用 StyleHandler 的样式到段落
        
        Args:
            paragraph: Word段落对象
            element_style: 元素样式
            is_heading: 是否为标题，标题不应用首行缩进
        """
        from docx.enum.text import WD_LINE_SPACING
        
        # 应用字体样式
        for run in paragraph.runs:
            font = run.font
            font.name = element_style.font.name
            font.size = Pt(element_style.font.size)
            font.bold = element_style.font.bold
            font.italic = element_style.font.italic
            
            if element_style.font.color:
                color_hex = element_style.font.color.replace('#', '')
                if len(color_hex) == 6:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    font.color.rgb = RGBColor(r, g, b)
        
        # 应用段落样式
        paragraph_format = paragraph.paragraph_format
        
        # 对齐方式
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        paragraph.alignment = alignment_map.get(element_style.paragraph.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # 行距处理：如果 line_spacing >= 20，认为是固定值（磅），否则认为是倍数
        if element_style.paragraph.line_spacing >= 20:
            # 固定值行距（磅）
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph_format.line_spacing = Pt(element_style.paragraph.line_spacing)
        else:
            # 倍数行距
            paragraph_format.line_spacing = element_style.paragraph.line_spacing
        
        # 段前段后间距
        paragraph_format.space_before = Pt(element_style.paragraph.space_before)
        paragraph_format.space_after = Pt(element_style.paragraph.space_after)
        
        # 缩进
        # 注意：配置中的值单位是厘米，需要转换为英寸
        # 1厘米 = 1/2.54 英寸
        paragraph_format.left_indent = Inches(element_style.paragraph.left_indent / 2.54)
        paragraph_format.right_indent = Inches(element_style.paragraph.right_indent / 2.54)
        
        # 首行缩进：标题不应用首行缩进，只有正文段落才应用
        if is_heading:
            # 标题不应用首行缩进
            paragraph_format.first_line_indent = Inches(0)
        elif element_style.paragraph.first_line_indent == 0 and element_style.font.size:
            # 正文段落：如果配置值为0，则根据字体大小动态计算两个字符的宽度
            # 两个字符宽度 = 字体大小 × 2（磅），转换为英寸
            # 1磅 = 1/72 英寸
            first_line_indent_inches = (element_style.font.size * 2) / 72.0
            paragraph_format.first_line_indent = Inches(first_line_indent_inches)
        else:
            # 使用配置值（厘米转英寸）
            paragraph_format.first_line_indent = Inches(element_style.paragraph.first_line_indent / 2.54)
        
        # 分页控制
        paragraph_format.keep_together = element_style.paragraph.keep_together
        paragraph_format.keep_with_next = element_style.paragraph.keep_with_next
        paragraph_format.page_break_before = element_style.paragraph.page_break_before
    
    def _apply_style(self, paragraph, style: WordStyle):
        """应用样式到段落"""
        from docx.enum.text import WD_LINE_SPACING
        
        # 设置字体
        for run in paragraph.runs:
            font = run.font
            font.name = style.font_name
            font.size = Pt(style.font_size)
            font.bold = style.bold
            font.italic = style.italic
            
            if style.color:
                font.color.rgb = RGBColor.from_string(style.color.replace('#', ''))
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        
        # 行距处理：如果 line_spacing >= 20，认为是固定值（磅），否则认为是倍数
        if style.line_spacing >= 20:
            # 固定值行距（磅）
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph_format.line_spacing = Pt(style.line_spacing)
        else:
            # 倍数行距
            paragraph_format.line_spacing = style.line_spacing
        
        paragraph_format.space_before = Pt(style.paragraph_spacing_before)
        paragraph_format.space_after = Pt(style.paragraph_spacing_after)
        
        # 设置对齐方式
        if style.alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style.alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif style.alignment == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _apply_code_style(self, run):
        """应用代码样式"""
        font = run.font
        font.name = 'Consolas'
        font.size = Pt(10)
        font.color.rgb = RGBColor(51, 51, 51)
    
    def _add_hyperlink(self, paragraph, url: str, text: str):
        """添加超链接"""
        # 创建超链接
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # 设置超链接样式
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        
        paragraph._p.append(hyperlink)
    
    def _set_paragraph_background(self, paragraph, color: str):
        """设置段落背景色"""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.replace("#", "")}" />')
        paragraph._p.get_or_add_pPr().append(shading_elm)
    
    def _set_paragraph_border(self, paragraph, side: str = 'left'):
        """设置段落边框"""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), 'CCCCCC')
        
        pBdr.append(border)
        pPr.append(pBdr)
    
    def _add_title(self, title: str):
        """添加文档标题"""
        title_paragraph = self.document.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_page_break(self):
        """添加分页符"""
        paragraph = self.document.add_paragraph()
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
    
    def set_page_margins(self, top: float = 1.0, bottom: float = 1.0, 
                        left: float = 1.0, right: float = 1.0):
        """设置页面边距（英寸）"""
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)
    
    def _prepare_charts(self, markdown_text: str):
        """准备图表：识别数据并生成图片
        
        Args:
            markdown_text: 原始Markdown文本
        """
        if not CHARTS_AVAILABLE:
            print("图表功能不可用：图表模块导入失败")
            print("请检查是否安装了 openai 和 matplotlib: pip install openai matplotlib")
            return
        
        print(f"开始准备图表，markdown文本长度: {len(markdown_text)}")
        
        try:
            # 初始化图表识别器和生成器
            print("初始化图表识别器...")
            recognizer = ChartRecognizer()
            print("初始化图表生成器...")
            self.chart_generator = ChartGenerator()
            
            # 解析图表数据
            print("解析图表数据...")
            self.chart_data = recognizer.recognize(self.chart_data_source)
            print(f"解析到 {len(self.chart_data)} 个图表")
            
            if not self.chart_data:
                print("未识别到图表数据")
                return
            
            # 生成所有图表图片
            for i, chart in enumerate(self.chart_data):
                try:
                    position = chart.get('position', '')
                    title = chart.get('title', '图表')
                    data = chart.get('data', {})
                    chart_type = chart.get('type', 'pie')
                    
                    print(f"生成图表 {i+1}/{len(self.chart_data)}: {title}, 类型: {chart_type}, position: {position}")
                    print(f"数据: {data}")
                    
                    # 根据图表类型生成不同的图表
                    if chart_type == 'bar':
                        # 生成柱状图
                        image_path = self.chart_generator.generate_bar_chart(
                            title=title,
                            data=data,
                            width_cm=self.config.get('chart_width', 14.0),
                            dpi=self.config.get('chart_dpi', 300)
                        )
                    elif chart_type == 'line':
                        # 生成折线图
                        image_path = self.chart_generator.generate_line_chart(
                            title=title,
                            data=data,
                            width_cm=self.config.get('chart_width', 14.0),
                            dpi=self.config.get('chart_dpi', 300)
                        )
                    else:
                        # 默认生成饼图
                        image_path = self.chart_generator.generate_pie_chart(
                            title=title,
                            data=data,
                            width_cm=self.config.get('chart_width', 14.0),
                            dpi=self.config.get('chart_dpi', 300)
                        )
                    
                    print(f"图表生成成功: {image_path}")
                    
                    # 存储图片路径，使用position作为key
                    self.chart_images[position] = image_path
                    self.temp_image_files.append(image_path)
                    
                except Exception as e:
                    print(f"生成图表失败 {chart.get('title', '')}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue
            
            print(f"图表准备完成，共生成 {len(self.chart_images)} 个图表")
            if self.chart_images:
                print(f"已生成的图表位置列表: {list(self.chart_images.keys())[:3]}...")  # 只打印前3个
                    
        except ValueError as e:
            print(f"图表识别器初始化失败（可能是API Key未设置）: {e}")
            self.chart_data = []
            self.chart_images = {}
        except Exception as e:
            print(f"准备图表失败: {e}")
            import traceback
            traceback.print_exc()
            self.chart_data = []
            self.chart_images = {}
    
    def _split_paragraph_by_punctuation(self, text: str) -> List[str]:
        """使用标点符号分割段落
        
        Args:
            text: 段落文本
            
        Returns:
            分割后的文本片段列表（每个片段包含标点符号）
        """
        import re
        # 标点符号：，。；：！？、
        # 使用正则表达式按标点符号分割，保留分隔符
        # 匹配标点符号（，。；：！？、）
        pattern = r'([，。；：！？、])'
        
        # 分割文本，保留分隔符
        parts = re.split(pattern, text)
        
        # 合并文本和标点符号
        segments = []
        i = 0
        while i < len(parts):
            if i < len(parts) - 1 and parts[i + 1] in ['，', '。', '；', '：', '！', '？', '、']:
                # 文本 + 标点符号
                segment = parts[i] + parts[i + 1]
                if segment.strip():
                    segments.append(segment.strip())
                i += 2
            else:
                # 只有文本，没有后续标点
                if parts[i].strip():
                    segments.append(parts[i].strip())
                i += 1
        
        # 如果没有分割出片段，返回原文本
        return segments if segments else [text]
    
    def _check_and_insert_chart(self, paragraph, paragraph_text: str):
        """检查段落文本，如果匹配position则插入图表
        
        Args:
            paragraph: Word段落对象
            paragraph_text: 段落文本内容
        """
        if not self.chart_images:
            return
        
        # 遍历所有图表，检查position是否匹配
        for position, image_path in list(self.chart_images.items()):
            # 检查是否是 after 或 before 模式
            is_after = position.startswith('after:')
            is_before = position.startswith('before:')
            
            if not (is_after or is_before):
                continue
            
            # 提取段落文本（完整段落）
            if is_after:
                paragraph_keyword = position.replace('after:', '').strip()
            else:
                paragraph_keyword = position.replace('before:', '').strip()
            
            # 匹配策略1：完整段落匹配（段落文本必须完全包含在paragraph_text中）
            matched = False
            match_mode = "full"  # "full" 或 "segment"
            
            if paragraph_keyword in paragraph_text:
                matched = True
                match_mode = "full"
                print(f"完整段落匹配成功: 段落文本长度={len(paragraph_keyword)}, 匹配位置='{paragraph_keyword[:80]}...'")
            else:
                # 匹配策略2：分段匹配（使用标点符号分割段落）
                print(f"完整段落匹配失败，尝试分段匹配...")
                segments = self._split_paragraph_by_punctuation(paragraph_text)
                print(f"段落分割为 {len(segments)} 个片段")
                
                # 检查关键词是否在任何一个片段中
                for segment in segments:
                    if paragraph_keyword in segment:
                        matched = True
                        match_mode = "segment"
                        print(f"分段匹配成功: 在片段中找到关键词，片段='{segment[:80]}...'")
                        break
                
                if not matched:
                    print(f"分段匹配也失败，跳过此图表: position='{position[:80]}...'")
            
            # 如果匹配成功，插入图表
            if matched:
                mode = "after" if is_after else "before"
                match_info = f"({mode}模式, {match_mode}匹配)"
                print(f"匹配到图表插入位置 {match_info}: 段落文本长度={len(paragraph_keyword)}, 匹配位置='{paragraph_keyword[:80]}...'")
                # 找到匹配，插入图表
                # 注意：无论完整匹配还是分段匹配，都插入到完整段落的前后，而不是片段位置
                try:
                    print(f"开始插入图表: {image_path}")
                    
                    # 检查文件是否存在
                    import os
                    if not os.path.exists(image_path):
                        print(f"警告: 图片文件不存在: {image_path}")
                        del self.chart_images[position]
                        continue
                    
                    # 获取文件大小
                    file_size = os.path.getsize(image_path)
                    print(f"图片文件大小: {file_size / 1024:.2f} KB")
                    
                    # 使用底层API插入段落
                    from docx.oxml import parse_xml
                    from docx.text.paragraph import Paragraph
                    
                    # 创建新的段落元素
                    p_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                    new_p = parse_xml(p_xml)
                    
                    if is_after:
                        # after模式：在当前段落后插入新段落
                        paragraph._element.addnext(new_p)
                    else:
                        # before模式：在当前段落前插入新段落
                        paragraph._element.addprevious(new_p)
                    
                    # 从新插入的元素创建段落对象
                    image_paragraph = Paragraph(new_p, paragraph._parent)
                    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    run = image_paragraph.add_run()
                    print(f"正在添加图片到Word文档（这可能需要几秒钟）...")
                    import time
                    start_time = time.time()
                    # 插入Word时使用配置的宽度（默认14.0厘米）
                    insert_width = self.config.get('chart_insert_width', 14.0)
                    run.add_picture(image_path, width=Cm(insert_width))
                    elapsed_time = time.time() - start_time
                    print(f"成功插入图表 ({mode}模式): {image_path} (宽度: {insert_width}cm, 耗时: {elapsed_time:.2f}秒)")
                    
                    # 可选：添加图表标题
                    if self.config.get('add_chart_title', False):
                        # 创建标题段落
                        title_p_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                        title_new_p = parse_xml(title_p_xml)
                        
                        if is_after:
                            # after模式：在图片段落后插入标题
                            new_p.addnext(title_new_p)
                        else:
                            # before模式：在图片段落前插入标题（标题在图片上方）
                            new_p.addprevious(title_new_p)
                        
                        title_paragraph = Paragraph(title_new_p, paragraph._parent)
                        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 从chart_data中找到对应的标题
                        chart_title = None
                        for chart in self.chart_data:
                            if chart.get('position') == position:
                                chart_title = chart.get('title', '')
                                break
                        if chart_title:
                            title_run = title_paragraph.add_run(f"图: {chart_title}")
                            title_run.font.size = Pt(10)
                    
                    # 从字典中移除，避免重复插入
                    del self.chart_images[position]
                    print(f"图表已插入并从待插入列表移除，剩余 {len(self.chart_images)} 个图表待插入")
                    
                except Exception as e:
                    print(f"插入图表失败 {position}: {e}")
                    import traceback
                    traceback.print_exc()
                    # 插入失败时不移除，保留在列表中，稍后在文档末尾尝试插入
                    print(f"图表插入失败，将保留在列表中，稍后在文档末尾尝试插入")
    
    def _insert_remaining_charts(self):
        """将未插入的图表插入到文档末尾"""
        if not self.chart_images:
            print("没有未插入的图表")
            return
        
        print(f"检测到 {len(self.chart_images)} 个图表未找到匹配段落，将插入到文档末尾")
        print(f"未插入的图表列表: {list(self.chart_images.keys())}")
        
        # 遍历所有未插入的图表
        for position, image_path in list(self.chart_images.items()):
            try:
                print(f"开始插入未匹配的图表到文档末尾: position={position[:100]}..., image_path={image_path}")
                
                # 检查文件是否存在
                import os
                if not os.path.exists(image_path):
                    print(f"错误: 图片文件不存在: {image_path}")
                    del self.chart_images[position]
                    continue
                
                # 获取文件大小
                file_size = os.path.getsize(image_path)
                print(f"图片文件大小: {file_size / 1024:.2f} KB")
                
                # 在文档末尾插入新段落，添加图片
                image_paragraph = self.document.add_paragraph()
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                run = image_paragraph.add_run()
                print(f"正在添加图片到Word文档末尾（这可能需要几秒钟）...")
                import time
                start_time = time.time()
                # 插入Word时使用配置的宽度（默认14.0厘米）
                insert_width = self.config.get('chart_insert_width', 14.0)
                run.add_picture(image_path, width=Cm(insert_width))
                elapsed_time = time.time() - start_time
                print(f"成功插入图表到文档末尾: {image_path} (宽度: {insert_width}cm, 耗时: {elapsed_time:.2f}秒)")
                
                # 可选：添加图表标题
                if self.config.get('add_chart_title', False):
                    title_para = self.document.add_paragraph()
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 从chart_data中找到对应的标题
                    chart_title = None
                    for chart in self.chart_data:
                        if chart.get('position') == position:
                            chart_title = chart.get('title', '')
                            break
                    if chart_title:
                        title_run = title_para.add_run(f"图: {chart_title}")
                        title_run.font.size = Pt(10)
                
                # 从字典中移除
                del self.chart_images[position]
                print(f"图表已插入到文档末尾并从待插入列表移除，剩余 {len(self.chart_images)} 个图表待插入")
                
            except Exception as e:
                print(f"插入图表到文档末尾失败 {position}: {e}")
                import traceback
                traceback.print_exc()
                # 即使失败也移除，避免重复尝试
                if position in self.chart_images:
                    del self.chart_images[position]
    
    def _cleanup_chart_images(self):
        """清理临时图片文件"""
        if not self.chart_generator:
            return
        
        for image_path in self.temp_image_files:
            try:
                self.chart_generator.cleanup(image_path)
            except Exception as e:
                print(f"清理图片文件失败 {image_path}: {e}")
        
        self.temp_image_files = []